import os
import xml.etree.ElementTree as ET
from docx import Document
import win32com.client

import zipfile

# Parent directory where archetype folders are located
parent_folder = "C:/Users/Asier/Desktop/TFG/mimic-openehr/archetype_oracle"

# Find the first folder or zip that starts with "archetype"
input_folder = None
for name in os.listdir(parent_folder):
    full_path = os.path.join(parent_folder, name)
    if name.startswith("archetype"):
        if os.path.isdir(full_path):
            input_folder = full_path
            break
        elif name.endswith(".zip"):
            # Extract zip to a temp folder
            extract_path = os.path.join(parent_folder, "archetypes")
            with zipfile.ZipFile(full_path, 'r') as zip_ref:
                zip_ref.extractall(extract_path)
            input_folder = extract_path
            break

if input_folder is None:
    raise FileNotFoundError("No folder or zip starting with 'archetype' found in the parent directory.")

output_folder = os.path.join(parent_folder, "documents")

# List all folders and subfolders
for dirpath, dirnames, filenames in os.walk(input_folder):
    print(f"Current directory: {dirpath}")
    # print(f"Subdirectories: {dirnames}")
    # print(f"Files: {filenames}")

# Extract namespace from root tag
def get_namespace(root):
    ns = {"ns": root.tag.split("}")[0].strip("{")}
    return ns

# Removes extra newlines and spaces from text
def clean_text(text):
    if text:
        return " ".join(text.strip().split())
    return ""

# Loop through XML files in folder and subfolders
for dirpath, _, filenames in os.walk(input_folder):
    
    # Create a new Word document for each subfolder
    subfolder_name = os.path.basename(dirpath)
    doc = Document()
    doc.add_heading(f"Archetype Extraction Report for {subfolder_name}", level=1)

    document_empty = True   # Flag to check if document is empty
    skipped_archetypes = 0  # Counter for skipped archetypes
   
    for filename in filenames:
        if filename.endswith(".xml"):
            xml_path = os.path.join(dirpath, filename)
            tree = ET.parse(xml_path)
            root = tree.getroot()
            ns = get_namespace(root)

            # Extract main elements
            archetype_id = root.find(".//ns:archetype_id/ns:value", ns).text
            lifecycle_state = root.find(".//ns:description/ns:lifecycle_state", ns).text
            category = root.find("./ns:definition/ns:rm_type_name", ns).text

            # Extract the part between the last two dots
            archetype_id_parts = archetype_id.split('.')
            if len(archetype_id_parts) > 2:
                name = archetype_id_parts[-2]

            # Extract languages
            languages = []
            for details in root.findall(".//ns:description/ns:details", ns):
                lang_code = details.find("./ns:language/ns:code_string", ns).text
                languages.append(lang_code)
                if lang_code == "en":
                    en_details = details

            # If no English version found, skip archetype
            if "en" not in languages:
                print(f"Warning: Skipping archetype '{archetype_id}' (No English version found)")
                skipped_archetypes += 1
                continue

            # Extract purpose, use, and misuse (with cleaning)
            purpose = clean_text(en_details.find("./ns:purpose", ns).text if en_details.find("./ns:purpose", ns) is not None else "")
            use = clean_text(en_details.find("./ns:use", ns).text if en_details.find("./ns:use", ns) is not None else "")
            misuse = clean_text(en_details.find("./ns:misuse", ns).text if en_details.find("./ns:misuse", ns) is not None else "")

            # Extract keywords
            keywords = [kw.text for kw in en_details.findall("./ns:keywords", ns) if kw.text]

            # Extract ontology/concepts
            concepts = []
            for en_term_def in root.findall(".//ns:ontology/ns:term_definitions[@language='en']/ns:items", ns):
                code = en_term_def.get("code")  # ex "at0001"
                text_elem = en_term_def.find("ns:items[@id='text']", ns)
                desc_elem = en_term_def.find("ns:items[@id='description']", ns)

                text = text_elem.text if text_elem is not None else ""
                desc = desc_elem.text if desc_elem is not None else ""

                concepts.append(f"{code}::{text} - {desc}")

            # Add data to the Word document
            doc.add_heading(name, level=2)
            doc.add_paragraph(f"\n**Archetype ID:** {archetype_id}")
            doc.add_paragraph(f"**Lifecycle State:** {lifecycle_state}")
            doc.add_paragraph(f"**Category:** {category}")
            doc.add_paragraph(f"**Languages:** {', '.join(languages)}")

            if purpose:
                doc.add_paragraph(f"**Purpose:** {purpose}")
            if use:
                doc.add_paragraph(f"**Use:** {use}")
            if misuse:
                doc.add_paragraph(f"**Misuse:** {misuse}")

            if keywords:
                doc.add_paragraph(f"**Keywords:** {', '.join(keywords)}")

            if concepts:
                doc.add_paragraph("**Concepts:**")
                for concept in concepts:
                    doc.add_paragraph(concept, style="List Bullet")

            doc.add_paragraph("\n")            
            document_empty = False  # Document has content



    # Save the document for the current subfolder if it is not empty
    if not document_empty:
        output_doc = os.path.join(output_folder, f"{subfolder_name}_report.docx")
        doc.save(output_doc)
        break  # Remove this line to process all subfolders

    
# Print the total number of skipped archetypes
print(f"Total number of skipped archetypes: {skipped_archetypes}")


def doc2pdf(folder_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Run Word in the background

    for file in os.listdir(folder_path):
        if file.endswith(".doc") or file.endswith(".docx"):
            doc_path = os.path.normpath(os.path.join(folder_path, file))
            pdf_path = os.path.normpath(os.path.join(folder_path, file.rsplit(".", 1)[0] + ".pdf"))

            doc = word.Documents.Open(doc_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 is the PDF format
            doc.Close()

    word.Quit()
    print("Conversion complete!")

# Convert the Word documents to PDFs
doc2pdf(output_folder)
